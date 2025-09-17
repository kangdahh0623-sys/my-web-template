# backend/app/services/report_generator.py
"""
급식 메뉴 최적화 결과 보고서 생성 서비스
가정통신문 PDF/워드 문서 생성
"""

import os
import logging
from datetime import datetime, timedelta
from typing import Dict, Any, List, Optional
from io import BytesIO
import pandas as pd

# PDF 생성용
try:
    from reportlab.lib.pagesizes import A4, letter
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch, cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

# 워드 문서 생성용
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from app.core.config import settings

logger = logging.getLogger(__name__)

class ReportGenerator:
    """급식 보고서 생성기"""
    
    def __init__(self):
        self.font_setup_done = False
        logger.info("ReportGenerator 초기화")
        
        if not REPORTLAB_AVAILABLE:
            logger.warning("ReportLab이 설치되지 않아 PDF 생성을 사용할 수 없습니다.")
        if not DOCX_AVAILABLE:
            logger.warning("python-docx가 설치되지 않아 워드 생성을 사용할 수 없습니다.")

    def setup_korean_font(self):
        """한글 폰트 설정 (PDF용)"""
        if not REPORTLAB_AVAILABLE or self.font_setup_done:
            return
            
        try:
            # 시스템 폰트 경로들 시도
            font_paths = [
                '/System/Library/Fonts/AppleGothic.ttf',  # macOS
                'C:/Windows/Fonts/malgun.ttf',             # Windows
                '/usr/share/fonts/truetype/nanum/NanumGothic.ttf',  # Ubuntu
                '/usr/share/fonts/nanum/NanumGothic.ttf'   # CentOS
            ]
            
            font_registered = False
            for font_path in font_paths:
                if os.path.exists(font_path):
                    pdfmetrics.registerFont(TTFont('NanumGothic', font_path))
                    font_registered = True
                    logger.info(f"한글 폰트 등록 성공: {font_path}")
                    break
            
            if not font_registered:
                logger.warning("한글 폰트를 찾을 수 없어 기본 폰트를 사용합니다.")
                
            self.font_setup_done = True
            
        except Exception as e:
            logger.error(f"한글 폰트 설정 실패: {e}")

    def generate_school_report(
        self, 
        menu_plan: pd.DataFrame, 
        summary: Dict[str, Any],
        school_info: Optional[Dict[str, str]] = None,
        format_type: str = "pdf"
    ) -> bytes:
        """
        학교 가정통신문 생성
        
        Args:
            menu_plan: 메뉴 계획 데이터프레임
            summary: 요약 정보
            school_info: 학교 정보
            format_type: "pdf" 또는 "docx"
            
        Returns:
            생성된 문서의 바이트 데이터
        """
        try:
            if format_type.lower() == "pdf":
                return self._generate_pdf_report(menu_plan, summary, school_info)
            elif format_type.lower() == "docx":
                return self._generate_docx_report(menu_plan, summary, school_info)
            else:
                raise ValueError(f"지원하지 않는 형식: {format_type}")
                
        except Exception as e:
            logger.error(f"보고서 생성 실패: {e}")
            raise

    def _generate_pdf_report(
        self, 
        menu_plan: pd.DataFrame, 
        summary: Dict[str, Any],
        school_info: Optional[Dict[str, str]]
    ) -> bytes:
        """PDF 보고서 생성"""
        if not REPORTLAB_AVAILABLE:
            raise RuntimeError("ReportLab이 설치되지 않았습니다.")
            
        self.setup_korean_font()
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            topMargin=1*cm,
            bottomMargin=1*cm,
            leftMargin=1.5*cm,
            rightMargin=1.5*cm
        )
        
        # 스타일 설정
        styles = getSampleStyleSheet()
        
        # 한글 스타일 정의
        korean_style = ParagraphStyle(
            'Korean',
            parent=styles['Normal'],
            fontName='NanumGothic' if self.font_setup_done else 'Helvetica',
            fontSize=10,
            leading=14
        )
        
        title_style = ParagraphStyle(
            'KoreanTitle',
            parent=korean_style,
            fontSize=16,
            spaceAfter=20,
            alignment=1,  # CENTER
            textColor=colors.darkblue
        )
        
        subtitle_style = ParagraphStyle(
            'KoreanSubtitle',
            parent=korean_style,
            fontSize=12,
            spaceBefore=15,
            spaceAfter=10,
            textColor=colors.darkgreen
        )
        
        # 문서 내용 구성
        content = []
        
        # 1. 제목
        school_name = school_info.get('name', '○○학교') if school_info else '○○학교'
        title = Paragraph(f"{school_name} 급식 계획서", title_style)
        content.append(title)
        content.append(Spacer(1, 0.3*cm))
        
        # 2. 발행 정보
        today = datetime.now()
        issue_date = today.strftime("%Y년 %m월 %d일")
        period_start = today + timedelta(days=7)  # 다음주부터
        period_end = period_start + timedelta(days=summary.get('days', 20))
        
        info_text = f"""
        <b>발행일:</b> {issue_date}<br/>
        <b>급식 기간:</b> {period_start.strftime('%Y년 %m월 %d일')} ~ {period_end.strftime('%Y년 %m월 %d일')}<br/>
        <b>총 급식일:</b> {summary.get('days', 20)}일<br/>
        <b>1인당 예산:</b> {summary.get('budget_won', 5370):,}원
        """
        content.append(Paragraph(info_text, korean_style))
        content.append(Spacer(1, 0.5*cm))
        
        # 3. 영양 요약
        content.append(Paragraph("영양 성분 요약", subtitle_style))
        nutrition_text = f"""
        • 1일 평균 칼로리: {summary.get('avg_kcal', 900):.0f} kcal<br/>
        • 총 급식비용: {summary.get('total_cost', 0):,.0f}원<br/>
        • 영양 균형: {'우수' if summary.get('feasible', True) else '조정 필요'}
        """
        content.append(Paragraph(nutrition_text, korean_style))
        content.append(Spacer(1, 0.3*cm))
        
        # 4. 메뉴표 (페이지당 10일씩)
        content.append(Paragraph("주간 급식 메뉴표", subtitle_style))
        
        # 합계 행 제외하고 실제 메뉴만
        menu_data = menu_plan[menu_plan['day'] != '합계'].copy()
        
        # 테이블 데이터 준비
        table_data = [['일차', '밥', '국', '반찬1', '반찬2', '반찬3', '간식']]
        
        for _, row in menu_data.iterrows():
            table_row = [
                f"DAY {row['day']}",
                str(row.get('rice', '')),
                str(row.get('soup', '')),
                str(row.get('side1', '')),
                str(row.get('side2', '')),
                str(row.get('side3', '')),
                str(row.get('snack', ''))
            ]
            table_data.append(table_row)
        
        # 테이블 스타일
        table = Table(table_data, colWidths=[1.5*cm, 2.5*cm, 2.5*cm, 2.5*cm, 2.5*cm, 2.5*cm, 2*cm])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightblue),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('FONTNAME', (0, 0), (-1, -1), 'NanumGothic' if self.font_setup_done else 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey])
        ]))
        
        content.append(table)
        content.append(Spacer(1, 0.5*cm))
        
        # 5. 알림사항
        content.append(Paragraph("학부모님께 알려드립니다", subtitle_style))
        notice_text = f"""
        1. 급식비는 매월 25일에 자동이체됩니다.<br/>
        2. 식단은 식재료 수급 상황에 따라 변경될 수 있습니다.<br/>
        3. 알레르기 유발 식품이 포함된 경우 별도 안내드립니다.<br/>
        4. 급식 관련 문의: 영양실 ({school_info.get('phone', '000-0000-0000') if school_info else '000-0000-0000'})<br/>
        <br/>
        항상 우리 아이들의 건강한 성장을 위해 최선을 다하겠습니다.<br/>
        <br/>
        <b>{school_name} 영양실</b>
        """
        content.append(Paragraph(notice_text, korean_style))
        
        # PDF 생성
        doc.build(content)
        buffer.seek(0)
        return buffer.getvalue()

    def _generate_docx_report(
        self, 
        menu_plan: pd.DataFrame, 
        summary: Dict[str, Any],
        school_info: Optional[Dict[str, str]]
    ) -> bytes:
        """워드 문서 보고서 생성"""
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx가 설치되지 않았습니다.")
        
        doc = Document()
        
        # 스타일 설정
        style = doc.styles['Normal']
        font = style.font
        font.name = '맑은 고딕'
        font.size = Pt(10)
        
        # 1. 제목
        school_name = school_info.get('name', '○○학교') if school_info else '○○학교'
        title = doc.add_heading(f'{school_name} 급식 계획서', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 2. 발행 정보
        today = datetime.now()
        issue_date = today.strftime("%Y년 %m월 %d일")
        period_start = today + timedelta(days=7)
        period_end = period_start + timedelta(days=summary.get('days', 20))
        
        info_para = doc.add_paragraph()
        info_para.add_run(f"발행일: {issue_date}\n").bold = True
        info_para.add_run(f"급식 기간: {period_start.strftime('%Y년 %m월 %d일')} ~ {period_end.strftime('%Y년 %m월 %d일')}\n")
        info_para.add_run(f"총 급식일: {summary.get('days', 20)}일\n")
        info_para.add_run(f"1인당 예산: {summary.get('budget_won', 5370):,}원\n")
        
        # 3. 영양 요약
        doc.add_heading('영양 성분 요약', level=1)
        nutrition_para = doc.add_paragraph()
        nutrition_para.add_run(f"• 1일 평균 칼로리: {summary.get('avg_kcal', 900):.0f} kcal\n")
        nutrition_para.add_run(f"• 총 급식비용: {summary.get('total_cost', 0):,.0f}원\n")
        nutrition_para.add_run(f"• 영양 균형: {'우수' if summary.get('feasible', True) else '조정 필요'}\n")
        
        # 4. 메뉴표
        doc.add_heading('주간 급식 메뉴표', level=1)
        
        # 합계 행 제외
        menu_data = menu_plan[menu_plan['day'] != '합계'].copy()
        
        # 테이블 생성
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Table Grid'
        
        # 헤더
        hdr_cells = table.rows[0].cells
        headers = ['일차', '밥', '국', '반찬1', '반찬2', '반찬3', '간식']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # 데이터 행들
        for _, row in menu_data.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = f"DAY {row['day']}"
            row_cells[1].text = str(row.get('rice', ''))
            row_cells[2].text = str(row.get('soup', ''))
            row_cells[3].text = str(row.get('side1', ''))
            row_cells[4].text = str(row.get('side2', ''))
            row_cells[5].text = str(row.get('side3', ''))
            row_cells[6].text = str(row.get('snack', ''))
        
        # 5. 알림사항
        doc.add_heading('학부모님께 알려드립니다', level=1)
        notice_para = doc.add_paragraph()
        notice_para.add_run("1. 급식비는 매월 25일에 자동이체됩니다.\n")
        notice_para.add_run("2. 식단은 식재료 수급 상황에 따라 변경될 수 있습니다.\n")
        notice_para.add_run("3. 알레르기 유발 식품이 포함된 경우 별도 안내드립니다.\n")
        notice_para.add_run(f"4. 급식 관련 문의: 영양실 ({school_info.get('phone', '000-0000-0000') if school_info else '000-0000-0000'})\n\n")
        notice_para.add_run("항상 우리 아이들의 건강한 성장을 위해 최선을 다하겠습니다.\n\n")
        
        signature = doc.add_paragraph()
        signature.add_run(f"{school_name} 영양실").bold = True
        signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # 바이트 스트림으로 저장
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

# 전역 서비스 인스턴스
report_generator = ReportGenerator()